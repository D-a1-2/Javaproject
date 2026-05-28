from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "202506050116-董政旭-AI使用辅助报告.docx"


def set_font(run, name="微软雅黑", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color="2E74B5")


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        shade(t.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=9.5, bold=(row == t.rows[0]))
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI 使用辅助报告")
    set_font(r, size=20, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("学号：202506050116    姓名：董政旭    课程项目：Java CLI 多策略爬虫")
    set_font(r, size=10.5)

    heading(doc, "一、AI 辅助使用概述")
    body(doc, "本项目在开发过程中使用 AI 作为辅助工具，主要用于理解课程要求、重构原有 Java 爬虫代码、设计项目结构、补充异常处理、生成实验报告初稿和检查编译运行结果。AI 仅作为辅助分析与代码建议工具，最终项目功能、文件结构和报告内容均经过人工确认与运行验证。")

    heading(doc, "二、使用 AI 的具体环节")
    table(
        doc,
        ["阶段", "AI 辅助内容", "本人完成或确认内容"],
        [
            ["需求分析", "从课程项目要求 PDF 中提取关键要求：CLI、MVC、Command、Strategy、异常体系、3 个以上网站、文件保存、git project 文件夹", "确认这些要求作为项目验收清单"],
            ["代码重构", "建议将原始单文件爬虫拆分为 controller、view、model、command、strategy、exception、service 包", "确认包结构，并检查代码是否符合课程知识点"],
            ["功能实现", "辅助生成 Command 接口、多个命令类、CrawlStrategy 接口、网页和股票爬取策略、结果保存类", "运行 javac 编译并执行 batch 命令验证功能"],
            ["异常处理", "辅助设计 NetworkException、ParseException、ValidationException、StorageException 等异常类", "确认 Controller 能统一捕获异常并显示友好提示"],
            ["报告撰写", "根据项目代码生成实验报告和 AI 使用辅助报告初稿", "核对姓名、学号、项目说明、测试结果和总结内容"],
        ],
    )

    heading(doc, "三、关键提示词记录")
    bullet(doc, "“帮我完成 Java 程序设计的期末课程任务，我的 VSCode 上已经有一个爬虫程序了，但可能不完全符合课程的项目要求，你帮我改进一下并完成项目报告，并且按要求上传到我的 git 仓库。”")
    bullet(doc, "“我的姓名是董政旭，学号 202506050116，另外还需要上传个 AI 使用辅助的报告。”")
    bullet(doc, "开发过程中还通过 AI 对 PDF 要求、现有 Java 代码、编译结果和运行输出进行分析。")

    heading(doc, "四、AI 输出的采纳与修改")
    body(doc, "AI 提供的代码并非直接无条件采纳。开发过程中根据课程要求进行了筛选：保留符合 Java 基础课程特点的 JDK 标准库实现，不引入复杂第三方依赖；将批量爬取站点调整为更稳定的公开网页，避免网络环境导致验收时失败；保留新浪财经股票策略作为扩展命令，体现原始爬虫主题。")

    heading(doc, "五、风险控制与诚信说明")
    bullet(doc, "对 AI 生成代码进行了本地编译测试，确认 `javac -encoding UTF-8` 能通过。")
    bullet(doc, "对 CLI 的 `batch` 命令进行了实际运行测试，确认能够爬取多个网站并保存到 `data` 目录。")
    bullet(doc, "报告中对 AI 的使用范围进行了说明，未隐瞒 AI 在需求整理、代码重构和文档生成中的辅助作用。")
    bullet(doc, "最终提交内容由本人确认，项目设计思路、测试结果和课程要求对应关系均可解释。")

    heading(doc, "六、总结")
    body(doc, "AI 在本项目中主要提升了需求梳理、代码组织和文档整理效率。通过 AI 辅助，原始爬虫程序被重构为更符合课程要求的 Java CLI 项目；同时，本人通过阅读和运行代码，加深了对 MVC、Command 模式、策略模式和异常体系的理解。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
