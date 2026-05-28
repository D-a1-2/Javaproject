from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "202506050116-董政旭-期末实验报告.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def set_run_font(run, name="微软雅黑", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color="2E74B5" if level == 1 else "1F4D78")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "E8EEF5")
        if widths:
            set_cell_width(hdr[i], widths[i])
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            if widths:
                set_cell_width(cells[i], widths[i])
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)


def main():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《高级程序设计》期末实验报告")
    set_run_font(run, size=20, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Java CLI 多策略爬虫项目开发全过程记录")
    set_run_font(run, size=12, color="555555")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("学号：202506050116    姓名：董政旭    项目名称：Java 期末爬虫项目")
    set_run_font(run, size=10.5)

    add_heading(doc, "一、项目目标", 1)
    add_body(doc, "本项目目标是完成一个可在命令行运行的 Java 爬虫程序。程序支持用户通过 CLI 输入命令，选择不同爬取策略，对普通网页、网页标题和新浪财经股票接口进行爬取，并将结果保存到本地文件。")
    add_table(
        doc,
        ["功能", "描述", "优先级"],
        [
            ["CLI 交互", "通过 help、crawl、batch、history、exit 命令完成操作", "高"],
            ["MVC 架构", "Controller 负责调度，View 负责显示，Model 保存爬取任务和结果", "高"],
            ["Command 模式", "每条命令封装为独立 Command 类，便于扩展", "高"],
            ["策略模式", "不同网站或数据类型由不同 CrawlStrategy 实现", "高"],
            ["异常体系", "定义网络、解析、校验、存储等自定义异常", "高"],
            ["文件保存", "每次爬取结果保存为 UTF-8 文本文件", "高"],
            ["批量爬取", "batch 命令默认爬取 3 个以上网站", "高"],
        ],
        widths=[1800, 5700, 1200],
    )

    add_heading(doc, "二、项目进展", 1)
    add_table(
        doc,
        ["阶段", "本周任务", "所学知识", "遇到的困难与解决", "AI 帮助"],
        [
            ["W1", "整理原始爬虫代码，明确课程要求", "HTTP 请求、字符编码、文件输出", "原代码是单文件结构，不符合 MVC；先梳理模块边界", "帮助提取 PDF 要求并拆解检查项"],
            ["W2", "设计 CLI 命令和项目包结构", "MVC、Command 模式", "命令解析容易和业务逻辑混在一起；改为独立 Command 类", "提供包结构建议并生成基础代码"],
            ["W3", "实现多种爬取策略", "策略模式、多态、接口抽象", "不同网站返回编码和内容格式不同；封装统一 HTTP 父类", "协助设计 CrawlStrategy 接口"],
            ["W4", "完善异常体系、保存结果、测试运行", "自定义异常、健壮性处理", "新浪财经接口可能 403；批量任务改为稳定站点并保留股票手动策略", "协助定位失败原因并优化批量流程"],
        ],
        widths=[900, 2000, 1800, 2600, 1700],
    )

    add_heading(doc, "三、项目结构", 1)
    add_body(doc, "最终项目放在 git 仓库的 project 文件夹中，采用 Maven 标准目录结构。核心包结构如下：")
    add_code(
        doc,
        """project/
├── pom.xml
├── README.md
├── data/
└── src/main/java/com/crawler/
    ├── App.java
    ├── controller/CrawlerController.java
    ├── view/ConsoleView.java
    ├── model/CrawlTask.java, CrawlResult.java
    ├── command/Command.java, CrawlCommand.java, BatchCommand.java, ...
    ├── strategy/CrawlStrategy.java, HtmlTextStrategy.java, TitleStrategy.java, SinaStockStrategy.java
    ├── service/ResultStorage.java
    └── exception/CrawlerException.java, NetworkException.java, ParseException.java, ..."""
    )

    add_heading(doc, "四、设计模式说明", 1)
    add_bullet(doc, "MVC：CrawlerController 负责接收命令并调用业务逻辑；ConsoleView 负责输入输出；CrawlTask 和 CrawlResult 负责数据表达。")
    add_bullet(doc, "Command 模式：Command 接口定义 execute 方法，CrawlCommand、BatchCommand、HistoryCommand、HelpCommand、ExitCommand 分别对应不同命令。")
    add_bullet(doc, "策略模式：CrawlStrategy 定义统一爬取接口，HtmlTextStrategy、TitleStrategy、SinaStockStrategy 可按命令动态选择。")
    add_bullet(doc, "异常体系：CrawlerException 为根异常，派生 NetworkException、ParseException、ValidationException、StorageException，便于 Controller 统一处理错误。")

    add_heading(doc, "类图（文字版）", 2)
    add_code(
        doc,
        """App
 └─ CrawlerController
     ├─ ConsoleView
     ├─ Command
     │   ├─ CrawlCommand
     │   ├─ BatchCommand
     │   ├─ HelpCommand
     │   ├─ HistoryCommand
     │   └─ ExitCommand
     ├─ CrawlerStrategyFactory
     │   └─ CrawlStrategy
     │       ├─ HtmlTextStrategy
     │       ├─ TitleStrategy
     │       └─ SinaStockStrategy
     └─ ResultStorage

CrawlerException
 ├─ NetworkException
 ├─ ParseException
 ├─ ValidationException
 └─ StorageException"""
    )

    add_heading(doc, "五、成果展示", 1)
    add_body(doc, "程序启动后显示欢迎信息和命令提示符。输入 batch 后会自动爬取多个预设网站，并将结果保存到 data 目录。")
    add_code(
        doc,
        """> batch
爬取完成：Example Domain
方式：网页标题快速爬取
保存文件：.../data/20260528_173146_838_网页标题快速爬取_Example_Domain.txt

爬取完成：Example Domain
方式：普通网页正文爬取
保存文件：.../data/20260528_173147_066_普通网页正文爬取_Example_Domain.txt

爬取完成：IANA 保留域名说明
方式：普通网页正文爬取
保存文件：.../data/20260528_173147_933_普通网页正文爬取_IANA_保留域名说明.txt

爬取完成：RFC Editor
方式：网页标题快速爬取
保存文件：.../data/20260528_173159_641_网页标题快速爬取_RFC_Editor.txt
批量爬取结束，成功任务数：4"""
    )

    add_heading(doc, "功能测试", 2)
    add_table(
        doc,
        ["功能", "测试命令", "测试结果", "备注"],
        [
            ["查看帮助", "help", "通过", "显示所有可用命令"],
            ["批量爬取", "batch", "通过", "成功爬取 4 个任务，满足 3 个以上网站要求"],
            ["网页正文爬取", "crawl html https://www.example.com", "通过", "提取网页正文并保存"],
            ["网页标题爬取", "crawl title https://www.rfc-editor.org", "通过", "提取标题并保存"],
            ["股票数据爬取", "crawl stock sh600519 贵州茅台", "通过", "可获取股票价格、涨跌幅和成交量"],
            ["异常处理", "crawl unknown test", "通过", "显示自定义校验错误，不导致程序崩溃"],
        ],
        widths=[1600, 2900, 1000, 3600],
    )

    add_heading(doc, "六、总结", 1)
    add_body(doc, "本次项目在原有简单爬虫基础上，重构为一个结构清晰、可扩展的 Java CLI 爬虫程序。通过 MVC 分层，程序界面、控制流程和数据模型职责更加明确；通过 Command 模式，新增命令只需要添加新的命令类；通过策略模式，新增网站或数据来源只需要实现新的 CrawlStrategy；通过自定义异常体系，网络失败、解析失败和参数错误能够被统一处理。")
    add_body(doc, "最终程序已实现课程要求中的 CLI、MVC、Command 模式、策略模式、异常体系、爬取 3 个以上网站和数据保存到文件。后续还可以继续扩展 JSON 解析、数据库保存、图形界面和定时爬取功能。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
